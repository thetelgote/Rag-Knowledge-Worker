import glob
from pathlib import Path
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# PDF
from langchain_community.document_loaders import PyPDFLoader

# Word
from docx import Document as DocxDocument


def load_txt_md(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def load_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def load_pdf(file_path):
    loader = PyPDFLoader(file_path)
    pages = loader.load()
    return "\n".join([p.page_content for p in pages])


def load_docs():
    docs = []

    BASE_DIR = Path(__file__).resolve().parent.parent
    base_path = str(BASE_DIR / "knowledge-base")

    print("📂 Searching in:", base_path)

    patterns = [
        "**/*.md",
        "**/*.txt",
        "**/*.pdf",
        "**/*.docx"
    ]

    files = []
    for pattern in patterns:
        files.extend(glob.glob(f"{base_path}/{pattern}", recursive=True))

    print("📄 Files found:", files)

    for f in files:
        try:
            ext = Path(f).suffix.lower()

            if ext in [".md", ".txt"]:
                text = load_txt_md(f)

            elif ext == ".pdf":
                text = load_pdf(f)

            elif ext == ".docx":
                text = load_docx(f)

            else:
                continue

            docs.append(Document(
                page_content=text,
                metadata={
                    "source": f,
                    "filename": Path(f).name.lower()
                }
            ))

        except Exception as e:
            print(f"❌ Error loading {f}: {e}")

    return docs


def run_ingest():
    print("📂 Loading docs...")
    docs = load_docs()

    print("Total Docs:", len(docs))

    if len(docs) == 0:
        print("❌ No documents found")
        return

    print("✂️ Splitting into chunks...")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )

    docs = splitter.split_documents(docs)

    print("🧠 Creating embeddings...")

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

    print("📦 Creating FAISS index...")

    db = FAISS.from_documents(docs, embeddings)
    db.save_local("faiss_index")

    print("✅ Ingestion complete!")


if __name__ == "__main__":
    run_ingest()